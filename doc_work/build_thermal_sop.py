from __future__ import annotations

import math
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REFERENCE = Path("/Users/mandinu/Downloads/Experimental_Toolkit_Photosynthesis_SOP_final.docx")
OUTPUT = Path("/Users/mandinu/Downloads/project/Experimental_Toolkit_Thermal_Insulation_SOP_Illustrated.docx")
IMAGE_DIR = Path("/Users/mandinu/Downloads/WhatsApp Unknown 2026-07-22 at 20.32.54")

IMAGES = {
    "coefficients": IMAGE_DIR / "WhatsApp Image 2026-07-03 at 16.57.02.jpeg",
    "full_photo": IMAGE_DIR / "WhatsApp Image 2026-07-15 at 04.56.05 (1).jpeg",
    "bubble_photo": IMAGE_DIR / "WhatsApp Image 2026-07-15 at 04.56.06.jpeg",
    "mylar_photo": IMAGE_DIR / "WhatsApp Image 2026-07-15 at 04.56.06 (2).jpeg",
    "bath_photo": IMAGE_DIR / "WhatsApp Image 2026-07-15 at 04.56.07.jpeg",
    "bare_photo": IMAGE_DIR / "WhatsApp Image 2026-07-15 at 05.02.37.jpeg",
    "full_graph": IMAGE_DIR / "WhatsApp Image 2026-07-15 at 04.57.44 (1).jpeg",
    "bubble_graph": IMAGE_DIR / "WhatsApp Image 2026-07-15 at 04.59.21.jpeg",
    "mylar_graph": IMAGE_DIR / "WhatsApp Image 2026-07-15 at 05.00.50.jpeg",
    "bare_graph": IMAGE_DIR / "WhatsApp Image 2026-07-15 at 05.02.05 (1).jpeg",
}

BLUE = "0F4761"
LIGHT_BLUE = "DDEBF7"
PALE_BLUE = "EAF3F7"
PALE_ORANGE = "FCE4D6"
PALE_GREEN = "E2F0D9"
PALE_RED = "F4CCCC"
GRAY = "F2F2F2"
BULLET_NUM_ID = None
NUMBER_NUM_ID = None


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + edge
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_text(cell, text, *, bold=False, align=None, size=None, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return p


def clear_body_keep_section(document):
    body = document._element.body
    section_properties = body.sectPr
    for child in list(body):
        if child is not section_properties:
            body.remove(child)


def create_numbering_definition(document, *, bullet):
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    abstract_id = (max(abstract_ids) + 1) if abstract_ids else 0
    num_id = (max(num_ids) + 1) if num_ids else 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.append(p_pr)
    if bullet:
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Symbol")
        fonts.set(qn("w:hAnsi"), "Symbol")
        r_pr.append(fonts)
        lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num.append(abs_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


def add_body(document, text="", *, bold=False, italic=False, align=None, keep_next=False):
    p = document.add_paragraph(style="Normal")
    if align is not None:
        p.alignment = align
    p.paragraph_format.keep_with_next = keep_next
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def add_bullet(document, text):
    p = document.add_paragraph(style="Normal")
    apply_numbering(p, BULLET_NUM_ID)
    p.add_run(text)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_numbered(document, text):
    p = document.add_paragraph(style="Normal")
    apply_numbering(p, NUMBER_NUM_ID)
    p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_heading(document, text, level, *, page_break_before=False):
    p = document.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = page_break_before
    return p


def add_caption(document, text):
    p = document.add_paragraph(style="Figure Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_together = True
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    return p


def add_image(document, path, width_inches, caption, alt_text):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width_inches))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", caption)
    add_caption(document, caption)


def add_callout(document, label, text, fill=PALE_BLUE):
    p = document.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        edge = OxmlElement(f"w:{side}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "8" if side == "left" else "2")
        edge.set(qn("w:space"), "4")
        edge.set(qn("w:color"), BLUE)
        borders.append(edge)
    p_pr.append(borders)
    r = p.add_run(label + " ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    p.add_run(text)
    return p


def add_page_break(document):
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def style_header_row(row):
    for cell in row.cells:
        shade_cell(cell, LIGHT_BLUE)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor.from_string(BLUE)
    set_repeat_table_header(row)
    prevent_row_split(row)


def add_metadata_table(document):
    table = document.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2450, 6910])
    values = [
        ("Author", "Timon Karl-Heinz Schwarz"),
        ("Version number", "1"),
        ("Last edited on", "24-July-2026"),
    ]
    for row, (left, right) in zip(table.rows, values):
        set_cell_text(row.cells[0], left, bold=True)
        set_cell_text(row.cells[1], right)
        prevent_row_split(row)
    return table


def add_checklist_table(document, rows):
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Qty.", "Item", "Purpose / requirement"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_text(cell, text, bold=True)
    for qty, item, purpose in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], qty, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[1], item)
        set_cell_text(cells[2], purpose)
        prevent_row_split(table.rows[-1])
    set_table_geometry(table, [800, 3300, 5260])
    style_header_row(table.rows[0])
    return table


def add_config_table(document):
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Configuration", "Construction", "Main heat-transfer control", "Expected behavior"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_text(cell, text, bold=True, size=9.5)
    rows = [
        ("Bare core (control)", "No insulation; identical probe depth and container.", "Reference for natural convection and radiation.", "Fast cooling."),
        ("Bubble wrap only", "Two uniform layers; close the neck without crushing the bubbles.", "Trapped air reduces conduction and convection.", "Moderate heat retention."),
        ("Mylar only", "One reflective layer with the shiny face outward; secure the neck.", "Reduces radiation only; air leaks can dominate.", "Highly sensitive to gaps."),
        ("Full MLI", "Cotton inner layer + bubble wrap + reflective Mylar outer layer.", "Combines conduction, convection, and radiation control.", "Best heat retention when tightly sealed."),
    ]
    fills = [GRAY, PALE_BLUE, PALE_ORANGE, PALE_GREEN]
    for data, fill in zip(rows, fills):
        cells = table.add_row().cells
        for cell, text in zip(cells, data):
            set_cell_text(cell, text, size=9.3)
            shade_cell(cell, fill)
        prevent_row_split(table.rows[-1])
    set_table_geometry(table, [1800, 3000, 2500, 2060])
    style_header_row(table.rows[0])
    return table


def add_procedure_table(document):
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Step", "Action", "Notes", "Duration", "Check"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_text(cell, text, bold=True, size=9.5)
    rows = [
        ("1", "Prepare a stable, dry work area.", "Keep the hot-water station separate from laptops, power banks, and cables.", "2 min.", ""),
        ("2", "Measure the ambient temperature.", "Record room temperature once before the first trial and again if the room changes noticeably.", "1 min.", ""),
        ("3", "Label the four test conditions.", "Bare, Bubble, Mylar, and Full MLI. Use identical cores or reuse the same core after returning it to the same starting condition.", "2 min.", ""),
        ("4", "Build each insulation package.", "Keep material area, probe position, and neck closure consistent. Do not compress the bubble or cotton layers.", "8 min.", ""),
        ("5", "Prepare the hot-water bath.", "Teacher/adult only: target approximately 80 °C. Confirm with a separate bath thermometer.", "5 min.", ""),
        ("6", "Condition the thermal core.", "Heat the core until its central probe reading is stable and within 80 ± 2 °C.", "3–5 min.", ""),
        ("7", "Transfer the core to the dry test station.", "Use gloves or tongs. Wipe external water away. Keep the probe at the same depth for every condition.", "< 30 s", ""),
        ("8", "Start the timer and record T₀.", "Start timing immediately after transfer. If T₀ is outside the agreed tolerance, recondition and repeat.", "A few seconds", ""),
        ("9", "Record temperature every minute.", "Take 16 readings from 0 through 15 min. Do not touch, move, unwrap, or shade the sample during the run.", "15 min.", ""),
        ("10", "Repeat for all four conditions.", "Use the same room position, container, water mass, timing interval, and measurement method.", "As needed", ""),
        ("11", "Enter observations in the telemetry sheet.", "Calculate temperature drop, retained temperature, model variance, and—if used—the correlation score.", "5 min.", ""),
        ("12", "Cool, dry, and store the equipment.", "Do not remove insulation until the core is safe to handle. Empty water only at a sink.", "5 min.", ""),
    ]
    for step, action, notes, duration, check in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], step, align=WD_ALIGN_PARAGRAPH.CENTER, size=9.2)
        set_cell_text(cells[1], action, size=9.2)
        set_cell_text(cells[2], notes, size=9.2)
        set_cell_text(cells[3], duration, align=WD_ALIGN_PARAGRAPH.CENTER, size=9.2)
        set_cell_text(cells[4], check, align=WD_ALIGN_PARAGRAPH.CENTER, size=9.2)
        prevent_row_split(table.rows[-1])
    set_table_geometry(table, [620, 2750, 3970, 1220, 800])
    style_header_row(table.rows[0])
    return table


def add_results_summary(document):
    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["Condition", "T₀ (°C)", "T₁₅ (°C)", "Drop (°C)", "Heat retained*", "Dashboard score", "Grade"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_text(cell, text, bold=True, size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ("Full MLI", 77.6, 62.7, 14.9, 25.0, "93%", "A"),
        ("Bubble wrap", 82.6, 52.6, 30.0, 52.1, "83%", "B"),
        ("Bare control", 80.1, 35.2, 44.9, 74.7, "94%", "A"),
        ("Mylar only", 80.5, 33.3, 47.2, 78.0, "3%", "F"),
    ]
    fills = [PALE_GREEN, PALE_BLUE, GRAY, PALE_RED]
    for data, fill in zip(rows, fills):
        cells = table.add_row().cells
        for i, value in enumerate(data):
            text = f"{value:.1f}" if isinstance(value, float) else str(value)
            set_cell_text(cells[i], text, size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER if i != 0 else None)
            shade_cell(cells[i], fill)
        prevent_row_split(table.rows[-1])
    set_table_geometry(table, [1700, 1050, 1050, 1100, 1400, 1750, 1310])
    style_header_row(table.rows[0])
    add_body(document, "*Heat retained is calculated relative to a 20 °C ambient reference: (T₁₅−20)/(T₀−20) × 100. Values are rounded.", italic=True)
    return table


def add_observed_data_table(document):
    times = list(range(16))
    full = [77.6, 76.6, 76.1, 75.6, 75.1, 74.5, 73.9, 73.5, 72.1, 70.6, 69.7, 68.5, 67.4, 65.8, 64.1, 62.7]
    bubble = [82.6, 79.0, 75.2, 72.5, 70.0, 67.9, 65.6, 64.0, 62.3, 60.6, 58.9, 57.5, 56.1, 54.8, 53.7, 52.6]
    mylar = [80.5, 68.6, 61.8, 55.8, 51.3, 47.8, 44.9, 42.5, 40.6, 38.9, 38.1, 36.3, 35.5, 34.6, 33.9, 33.3]
    bare = [80.1, 73.9, 66.0, 59.9, 54.5, 50.1, 47.1, 44.6, 42.5, 40.9, 39.5, 38.2, 37.3, 36.5, 35.8, 35.2]
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Time (min)", "Full MLI (°C)", "Bubble (°C)", "Mylar (°C)", "Bare (°C)"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_text(cell, text, bold=True, size=8.8, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_values in zip(times, full, bubble, mylar, bare):
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            text = str(value) if isinstance(value, int) else f"{value:.1f}"
            set_cell_text(cell, text, size=8.6, align=WD_ALIGN_PARAGRAPH.CENTER)
        prevent_row_split(table.rows[-1])
    set_table_geometry(table, [1500, 1965, 1965, 1965, 1965])
    style_header_row(table.rows[0])
    return table


def add_equation_paragraph(document, label, equation):
    p = document.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(label + ": ")
    r.bold = True
    e = p.add_run(equation)
    e.font.name = "Cambria Math"
    return p


def build():
    global BULLET_NUM_ID, NUMBER_NUM_ID
    if not REFERENCE.exists():
        raise FileNotFoundError(REFERENCE)
    for key, path in IMAGES.items():
        if not path.exists():
            raise FileNotFoundError(f"{key}: {path}")

    shutil.copy2(REFERENCE, OUTPUT)
    document = Document(OUTPUT)
    clear_body_keep_section(document)
    BULLET_NUM_ID = create_numbering_definition(document, bullet=True)
    NUMBER_NUM_ID = create_numbering_definition(document, bullet=False)

    section = document.sections[0]
    section.different_first_page_header_footer = True

    document.core_properties.title = "Experimental Toolkit Procedure - Thermal Insulation and Cooling Rates"
    document.core_properties.subject = "Illustrated classroom SOP for comparing passive thermal insulation systems"
    document.core_properties.keywords = "thermal insulation, Newton's law of cooling, MLI, bubble wrap, Mylar, classroom experiment"

    # Tighten the source-derived Caption style for figure-heavy pages.
    if "Figure Caption" not in [s.name for s in document.styles]:
        caption = document.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
        caption.base_style = document.styles["Normal"]
    else:
        caption = document.styles["Figure Caption"]
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(BLUE)

    # Title and front matter
    add_heading(document, "Experimental Toolkit Procedure - Thermal Insulation and Cooling Rates for Spacecraft Thermal Protection", 1)
    add_heading(document, "GENERAL INFORMATION", 2)
    add_metadata_table(document)

    add_heading(document, "SHORT EXPERIMENT DESCRIPTION", 2)
    add_body(
        document,
        "This classroom experiment compares how four passive insulation configurations slow the cooling of an identical heated thermal core: a bare control, bubble wrap, a reflective Mylar layer, and a full multilayer insulation (MLI) package made from cotton, bubble wrap, and Mylar. Students collect one temperature reading per minute for 15 minutes, compare the measured curves with Newtonian cooling predictions, and connect conduction, convection, and radiation to spacecraft thermal-control design.",
    )
    add_body(document, "Recommended class format: four groups working in parallel. Expected duration: approximately 55–70 minutes, including setup, a 15-minute measurement run, and post-processing.")

    add_heading(document, "LEARNING OBJECTIVES", 2)
    for item in [
        "Distinguish conduction, convection, and thermal radiation in a practical system.",
        "Explain why trapped air and layered insulation can reduce heat transfer.",
        "Collect a controlled time series and compare observed and predicted cooling curves.",
        "Evaluate why material choice alone is insufficient when gaps, compression, or inconsistent probe placement introduce experimental error.",
    ]:
        add_bullet(document, item)

    add_heading(document, "HARDWARE CHECKLIST", 2, page_break_before=True)
    add_heading(document, "From the school / students:", 4)
    add_checklist_table(
        document,
        [
            ("1", "Heat-safe bucket or deep container", "Used as a water bath and secondary containment."),
            ("1", "Kettle or other supervised hot-water source", "Teacher/adult use only."),
            ("2+", "Digital probe thermometers", "One for the thermal core; one to verify the water bath."),
            ("1", "Timer or stopwatch", "Must display one-minute intervals."),
            ("1", "Heat-resistant gloves or tongs", "Required for handling the heated core."),
            ("—", "Towels and a stable dry surface", "Keep water away from electronic equipment."),
        ],
    )
    add_heading(document, "Included in the toolkit / prepared for each group:", 4)
    add_checklist_table(
        document,
        [
            ("4 or 1 reused", "Identical thermal core / metal capsule with probe access", "Use the same water mass, container, and probe depth for every condition."),
            ("1 set", "Cotton or soft fibrous wrap", "Inner conductive barrier for the full MLI package."),
            ("1 roll", "Bubble wrap", "Convective/conductive shield; do not crush the bubbles."),
            ("1 sheet", "Reflective Mylar or emergency blanket", "Radiation shield and outer layer of the full MLI package."),
            ("As needed", "Tape and elastic bands", "Secure layers without compressing them."),
            ("1 per student", "Worksheet / lab report", "Hypothesis, raw data, calculations, and evaluation."),
        ],
    )

    add_heading(document, "SAFETY AND CONTROL REQUIREMENTS", 2)
    add_callout(
        document,
        "HOT-WATER HAZARD.",
        "Water near 80 °C can cause serious scalds. An adult must prepare the bath and handle the heated core with gloves or tongs. Keep the bucket below waist height on a stable surface. Never use boiling water, never point a sealed hot container toward a person, and never immerse electrical displays, connectors, or power supplies.",
        fill=PALE_RED,
    )
    for item in [
        "Use only heat-safe, undamaged containers. Do not test pressurized or completely sealed containers unless they are designed for hot liquids.",
        "Keep all electronics and cables dry; only the intended metal probe may contact the water.",
        "Use identical starting temperature, water mass, probe depth, ambient location, and measurement interval for fair comparison.",
        "Stop the test if a probe slips, insulation becomes wet, a container leaks, or the starting temperature is outside the agreed tolerance.",
    ]:
        add_bullet(document, item)

    add_heading(document, "TEST CONFIGURATIONS", 2)
    add_config_table(document)
    add_callout(
        document,
        "Optional extension.",
        "A cotton-only configuration may be added as a fifth trial. The supplied teaching model lists cotton as an intermediate solution, while the attached physical telemetry focuses on bare, bubble, Mylar-only, and full MLI configurations.",
    )
    add_callout(
        document,
        "Fair-test rule.",
        "When starting temperatures differ, do not rank insulation by final temperature alone. Compare temperature drop, normalized heat retained above ambient, or a fitted value of k.",
    )

    add_heading(document, "ILLUSTRATED SETUP", 2, page_break_before=True)
    add_heading(document, "Prepare and verify the hot-water station", 3)
    add_image(
        document,
        IMAGES["bath_photo"],
        3.0,
        "Figure 1. Water-bath temperature is checked with a separate probe before conditioning the thermal core.",
        "Green bucket used as a hot-water bath with a digital probe thermometer showing the bath temperature.",
    )
    add_image(
        document,
        IMAGES["bare_photo"],
        3.0,
        "Figure 2. Bare thermal core in the bath. Keep the probe centered and at a repeatable depth.",
        "Bare circular thermal core immersed in a green water bath with a digital probe thermometer inserted at its center.",
    )

    add_heading(document, "BUILD THE INSULATION PACKAGES", 2, page_break_before=True)
    add_heading(document, "Bubble-wrap configuration", 3)
    add_body(document, "Wrap the core with two even layers of bubble wrap. Close the neck around the probe, but do not crush the air pockets. Use the same overlap and tape position throughout the trial.")
    add_image(
        document,
        IMAGES["bubble_photo"],
        2.45,
        "Figure 3. Bubble-wrap-only capsule with the probe secured at the neck.",
        "Thermal capsule wrapped in clear bubble wrap with a digital thermometer probe inserted through the top.",
    )
    add_heading(document, "Mylar-only configuration", 3)
    add_body(document, "Apply one complete reflective layer with the shiny surface outward. Seal the neck carefully. A loose opening allows warm air to escape and can overwhelm the benefit of the radiation shield.")
    add_image(
        document,
        IMAGES["mylar_photo"],
        2.45,
        "Figure 4. Mylar-only capsule. The reflective layer is visible around the full body.",
        "Thermal capsule wrapped in reflective silver Mylar with a digital thermometer at the top.",
    )

    add_heading(document, "FULL MULTILAYER INSULATION (MLI)", 2, page_break_before=True)
    add_body(document, "Build the full package from the core outward: a uniform cotton layer, two layers of bubble wrap, and a reflective Mylar outer layer. Secure each layer independently so that the outer band does not compress the insulation. Close the neck around the probe to limit convective leakage.")
    add_numbered(document, "Place cotton evenly around the dry core; avoid thick folds.")
    add_numbered(document, "Add bubble wrap with consistent overlap and intact air cells.")
    add_numbered(document, "Finish with Mylar, shiny side outward, and close the neck.")
    add_numbered(document, "Check that the probe tip remains at the same central depth used for the other configurations.")
    add_image(
        document,
        IMAGES["full_photo"],
        3.0,
        "Figure 5. Multilayer test package during measurement; the reflective layer and sealed probe opening are visible.",
        "Insulated thermal package in a green bucket with reflective material and two digital probe thermometers visible.",
    )
    add_heading(document, "Model coefficients supplied with the teaching tool", 3)
    add_image(
        document,
        IMAGES["coefficients"],
        6.15,
        "Figure 6. Approximate Newtonian cooling coefficients used by the teaching model.",
        "Dark dashboard panel listing approximate cooling coefficients for bare core, cotton wrap, bubble wrap, and multilayer insulation.",
    )
    add_callout(
        document,
        "Interpret with care.",
        "The coefficient panel is a model reference, not a guaranteed result. Actual cooling depends on ambient temperature, core size, water mass, probe position, compression, drafts, and the quality of the neck seal.",
    )

    add_heading(document, "PROCEDURE", 2, page_break_before=True)
    add_body(document, "Complete the following steps for each insulation condition. Groups may work in parallel if all groups have identical equipment and the same ambient environment.")
    add_procedure_table(document)

    add_heading(document, "DATA RECORDING AND CALCULATION", 2)
    add_body(document, "Record temperature at 0, 1, 2, …, 15 minutes. Do not substitute readings taken at irregular times without recording the actual timestamps.")
    add_equation_paragraph(document, "Temperature drop", "ΔTdrop = T₀ − T₁₅")
    add_equation_paragraph(document, "Newtonian cooling model", "T(t) = Tamb + (T₀ − Tamb)e^(−kt)")
    add_equation_paragraph(document, "Estimated cooling coefficient", "k = −ln[(Tt − Tamb)/(T₀ − Tamb)] / t")
    add_body(document, "Use temperatures in degrees Celsius and time in minutes, so k is reported in min⁻¹. The logarithm expression is valid only while Tt remains above Tamb.")

    add_heading(document, "REFERENCE OBSERVATIONS FROM THE ATTACHED TEST", 2, page_break_before=True)
    add_body(document, "The following values were transcribed from the supplied 16-point telemetry screenshots. They are included as a worked example and should not replace student measurements.")
    add_results_summary(document)
    add_heading(document, "Complete observed temperature series", 3)
    add_observed_data_table(document)

    add_callout(
        document,
        "Worked-example interpretation.",
        "Full MLI retained the most heat (62.7 °C at 15 min), followed by bubble wrap (52.6 °C). The bare and Mylar-only samples reached 35.2 °C and 33.3 °C. The unusually weak Mylar-only result and 3% dashboard correlation indicate a construction or control issue—such as convective leakage—not a universal failure of reflective insulation.",
    )

    add_heading(document, "PREDICTION–OBSERVATION GRAPHS", 2, page_break_before=True)
    add_image(
        document,
        IMAGES["full_graph"],
        6.15,
        "Figure 7. Full MLI: 93% dashboard correlation and high stability.",
        "Graph comparing predicted and observed full multilayer insulation cooling curves over 15 minutes.",
    )
    add_image(
        document,
        IMAGES["bubble_graph"],
        6.15,
        "Figure 8. Bubble wrap: 83% dashboard correlation and moderate isolation.",
        "Graph comparing predicted and observed bubble-wrap cooling curves over 15 minutes.",
    )

    add_heading(document, "PREDICTION–OBSERVATION GRAPHS (CONTINUED)", 2, page_break_before=True)
    add_image(
        document,
        IMAGES["mylar_graph"],
        6.15,
        "Figure 9. Mylar-only layer: 3% dashboard correlation; the panel flags convective leakage.",
        "Graph showing rapid observed cooling for the Mylar-only capsule compared with a much slower predicted curve.",
    )
    add_image(
        document,
        IMAGES["bare_graph"],
        6.15,
        "Figure 10. Bare control: 94% dashboard correlation and high stability of the fitted model.",
        "Graph comparing predicted and observed cooling curves for the bare control capsule over 15 minutes.",
    )

    add_heading(document, "EXPERIMENT POST-PROCESSING", 2, page_break_before=True)
    add_heading(document, "Student analysis", 3)
    NUMBER_NUM_ID = create_numbering_definition(document, bullet=False)
    for item in [
        "Plot all four measured curves on the same axes using identical scales.",
        "Calculate ΔTdrop and estimate k for each configuration.",
        "Rank the configurations using a normalized measure that accounts for different starting temperatures.",
        "Identify at least two possible sources of uncertainty and explain the direction in which each could change the result.",
        "Compare the physical outcome with the model, then explain any disagreement using heat-transfer mechanisms rather than simply calling it an error.",
    ]:
        add_numbered(document, item)

    add_heading(document, "Class discussion prompts", 3)
    for item in [
        "Why can a reflective layer work well in vacuum but perform poorly in air when warm air can escape?",
        "Why should bubble wrap and cotton not be tightly compressed?",
        "Which variables must be controlled before comparing two cooling coefficients?",
        "How is spacecraft multilayer insulation similar to—and different from—the classroom package?",
    ]:
        add_bullet(document, item)

    add_heading(document, "Acceptance checklist", 3)
    checklist = document.add_table(rows=1, cols=2)
    checklist.style = "Table Grid"
    set_cell_text(checklist.rows[0].cells[0], "Check", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(checklist.rows[0].cells[1], "Completion criterion", bold=True)
    criteria = [
        "All four conditions used the same core, water mass, probe depth, and recording interval.",
        "Every trial includes 16 time–temperature pairs from 0 through 15 minutes.",
        "Ambient temperature and starting-temperature tolerance are documented.",
        "Graphs have units, labels, a legend, and an equal time axis.",
        "Conclusions distinguish material performance from build quality and experimental uncertainty.",
        "All hot equipment has cooled, been dried, and been stored safely.",
    ]
    for criterion in criteria:
        cells = checklist.add_row().cells
        set_cell_text(cells[0], "", align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[1], criterion)
        prevent_row_split(checklist.rows[-1])
    set_table_geometry(checklist, [900, 8460])
    style_header_row(checklist.rows[0])

    add_callout(
        document,
        "Teacher note.",
        "This experiment demonstrates comparative thermal behavior under classroom conditions. It does not certify any material for aerospace, pressure-vessel, fire-protection, or personal-protective use.",
        fill=PALE_ORANGE,
    )

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
